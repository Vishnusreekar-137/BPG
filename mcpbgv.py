import os
import json
import asyncio
import threading
import argparse
from typing import Dict, Any, Optional, Tuple, List
from dotenv import load_dotenv
from langchain_anthropic import ChatAnthropic
from langchain_together import ChatTogether
from langchain.prompts import PromptTemplate
from langchain_core.runnables import RunnableSequence
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.platypus import Paragraph, Table, Spacer
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches
from mcp.server.fastmcp import FastMCP
from fastapi import FastAPI, Request
from fastapi import Body
from pydantic import BaseModel
import uvicorn
import logging
import httpx
import requests
import re
import traceback
import sys

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger("BPG-Analyzer")

# Load environment variables
load_dotenv()

# Validate environment variables
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY")
TOGETHER_API_KEY = os.getenv("TOGETHER_API_KEY")

if not CLAUDE_API_KEY or not TOGETHER_API_KEY:
    raise ValueError("CLAUDE_API_KEY and TOGETHER_API_KEY environment variables must be set")

# Initialize LLMs
claude_llm = ChatAnthropic(
    model_name="claude-3-haiku-20240307",
    anthropic_api_key=CLAUDE_API_KEY,
    temperature=0.1,
    max_tokens=1000
)

together_llm = ChatTogether(
    model="mistralai/Mixtral-8x7B-Instruct-v0.1",
    together_api_key=TOGETHER_API_KEY,
    temperature=0.2
)

# Initialize FastAPI app and MCP server
app = FastAPI(title="Business Process Guide Generator")
mcp = FastMCP("bpg-generator")

# Pydantic model for user story input
class UserStoryInput(BaseModel):
    Title: str
    Description: str
    AcceptanceCriteria: List[str]
    AdditionalInformation: Optional[str] = None

# Lifespan event handlers (replacing deprecated on_event)
@app.on_event("startup")
async def startup_event():
    logger.info("Server starting up...")
    routes = [route.path for route in app.routes]
    logger.info(f"Registered routes: {routes}")

@app.on_event("shutdown")
async def shutdown_event():
    logger.info("Server shutting down...")

# -------------------------------
# Step 1a: Generate Mermaid with Claude
# -------------------------------

def generate_mermaid(description: str, theme: str = "forest") -> str:
    try:
        prompt_template = PromptTemplate(
            input_variables=["description"],
            template="Generate only the Mermaid syntax for a flowchart representing this process: {description}. Return ONLY the syntax starting with 'graph', with no additional text, comments, or formatting. Use 'finish' instead of 'end'."
        )

        chain = RunnableSequence(prompt_template | claude_llm)
        raw_content = chain.invoke({"description": description}).content.strip()

        def clean_mermaid_code(content: str) -> str:
            lines = content.split('\n')
            valid_lines = []
            for line in lines:
                line = line.strip()
                if line and (line.startswith("graph") or line.startswith("flowchart") or valid_lines):
                    if 'end' in line.lower():
                        line = line.replace('end', 'finish').replace('End', 'Finish')
                    valid_lines.append(line)
            cleaned = '\n'.join(valid_lines)
            return cleaned if cleaned.startswith(("graph", "flowchart")) else None

        mermaid_code = clean_mermaid_code(raw_content)
        if not mermaid_code:
            logger.warning("Generated content is not valid Mermaid syntax.")
            return None

        init_block = f"%%{{init: {{'theme': '{theme}'}}}}%%"
        if "%%{init:" not in mermaid_code:
            return f"{init_block}\n{mermaid_code}"
        return mermaid_code
    except Exception as e:
        logger.error(f"Error generating Mermaid syntax: {e}")
        traceback.print_exc()
        return None

# -------------------------------
# Step 1: Generate Documentation
# -------------------------------

def generate_documentation(user_story: Dict[str, Any]) -> Dict[str, Any]:
    try:
        description = user_story.get("Description", "No description available")
        
        # 🧠 Step 1: Generate Mermaid Flowchart (Claude)
        mermaid_syntax = generate_mermaid(description)

        # 🧠 Step 2: Use Together to infer common defects
        defect_prompt = PromptTemplate(
            input_variables=["description"],
            template="""From the following feature description, identify potential edge cases or user-related defects that may occur during implementation or usage. Return ONLY a valid JSON array like:
[
  {{"defect": "...", "occurrence": "..."}},
  ...
]
Description:
{description}"""
        )
        defect_chain = RunnableSequence(defect_prompt | together_llm)
        defects_json = defect_chain.invoke({"description": description}).content.strip()
        # Robust JSON parsing
        common_defects = None
        try:
            common_defects = json.loads(defects_json)
            if not isinstance(common_defects, list):
                raise ValueError("Defects must be a JSON array")
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON for defects: {defects_json} (Error: {e})")
            json_match = re.search(r'\[.*\]', defects_json, re.DOTALL)
            if json_match:
                try:
                    common_defects = json.loads(json_match.group(0))
                    if not isinstance(common_defects, list):
                        raise ValueError("Extracted defects must be a JSON array")
                    logger.warning(f"Recovered valid JSON for defects: {common_defects}")
                except json.JSONDecodeError as e2:
                    logger.error(f"Failed to recover valid JSON for defects: {e2}")
                    common_defects = []
            else:
                common_defects = []

        # 🧠 Step 3: Use Together to infer Appendix table data
        appendix_prompt = PromptTemplate(
            input_variables=["description"],
            template="""Based on the feature below, extract relevant configuration parameters or rules that developers or testers should be aware of. Return ONLY a valid JSON array like:
[
  {{"parameter": "...", "value": "..."}},
  ...
]
Description:
{description}"""
        )
        appendix_chain = RunnableSequence(appendix_prompt | together_llm)
        appendix_json = appendix_chain.invoke({"description": description}).content.strip()
        # Robust JSON parsing
        appendix_table = None
        try:
            appendix_table = json.loads(appendix_json)
            if not isinstance(appendix_table, list):
                raise ValueError("Appendix table must be a JSON array")
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON for appendix: {appendix_json} (Error: {e})")
            json_match = re.search(r'\[.*\]', appendix_json, re.DOTALL)
            if json_match:
                try:
                    appendix_table = json.loads(json_match.group(0))
                    if not isinstance(appendix_table, list):
                        raise ValueError("Extracted appendix must be a JSON array")
                    logger.warning(f"Recovered valid JSON for appendix: {appendix_table}")
                except json.JSONDecodeError as e2:
                    logger.error(f"Failed to recover valid JSON for appendix: {e2}")
                    appendix_table = []
            else:
                appendix_table = []

        return {
            "description": description,
            "flowchart": mermaid_syntax,
            "common_defects": common_defects,
            "appendix_table": appendix_table
        }

    except Exception as e:
        logger.error(f"Error in generate_documentation: {e}")
        traceback.print_exc()
        return None

# -------------------------------
# Step 2: Render Mermaid PNG via Kroki
# -------------------------------

def render_flowchart_to_files(mermaid_code: str, png_file: str = "flowchart.png", timeout: int = 30) -> Optional[str]:
    if not mermaid_code:
        logger.error("No Mermaid code provided.")
        return None

    cleaned_code = '\n'.join(line for line in mermaid_code.split('\n') if line.strip() and not line.strip().startswith('---'))

    try:
        response = requests.post(
            "https://kroki.io/mermaid/png",
            data=cleaned_code.encode("utf-8"),
            headers={"Content-Type": "text/plain; charset=utf-8"},
            timeout=timeout
        )
        response.raise_for_status()

        with open(png_file, "wb") as f:
            f.write(response.content)
        logger.info(f"✅ Flowchart PNG saved to: {png_file}")
        return png_file
    except Exception as e:
        logger.error(f"Error rendering PNG: {e}")
        traceback.print_exc()
        return None

# -------------------------------
# Step 3: Create PDF
# -------------------------------

def create_bpg_pdf(documentation: Dict, image_path: str, output_path: str = "business_process_guide.pdf") -> Optional[str]:
    try:
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        normal = styles['Normal']
        heading = styles['Heading2']
        elements = []

        # Title
        elements.append(Paragraph("Business Process Guide", styles['Heading1']))
        elements.append(Spacer(1, 0.2 * inch))

        # Description
        elements.append(Paragraph("Description:", heading))
        elements.append(Paragraph(documentation.get("description", "No description available"), normal))
        elements.append(Spacer(1, 0.15 * inch))

        # Process Flow
        elements.append(Paragraph("Process Flow:", heading))
        if image_path and os.path.exists(image_path):
            img = Image(image_path, width=4.5*inch, height=5.7*inch)
            elements.append(img)
        elements.append(PageBreak())

        # Common Defects Table
        elements.append(Paragraph("Common Defects:", heading))
        if documentation.get("common_defects"):
            defect_data = [["Defect", "Occurrence"]]
            for d in documentation["common_defects"]:
                defect_data.append([Paragraph(d.get("defect", ""), normal), d.get("occurrence", "")])
            defect_table = Table(defect_data, colWidths=[2.2*inch, 4.2*inch])
            defect_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey)
            ]))
            elements.append(defect_table)
            elements.append(Spacer(1, 0.2 * inch))

        # Appendix Table
        elements.append(Paragraph("Appendix:", heading))
        if documentation.get("appendix_table"):
            appendix_data = [["Parameter", "Value"]]
            for p in documentation["appendix_table"]:
                appendix_data.append([
                    Paragraph(p.get("parameter", ""), normal),
                    Paragraph(p.get("value", ""), normal)
                ])
            appendix_table = Table(appendix_data, colWidths=[2.2*inch, 4.2*inch])
            appendix_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            elements.append(appendix_table)

        doc.build(elements)
        logger.info(f"✅ PDF saved to: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"PDF Error: {e}")
        traceback.print_exc()
        return None

# -------------------------------
# Step 4: Create Word
# -------------------------------

def create_bpg_word(documentation: Dict, png_path: str, output_path: str = "business_process_guide.docx") -> Optional[str]:
    try:
        if os.path.exists(output_path):
            with open(output_path, 'rb'):
                pass

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        doc.add_heading("Business Process Guide", 1)
        doc.add_heading("Description:", 2)
        doc.add_paragraph(documentation.get("description", "No description available"))

        doc.add_heading("Process Flow:", 2)
        if png_path and os.path.exists(png_path):
            doc.add_picture(png_path, width=Inches(6), height=Inches(6.5))

        doc.add_page_break()

        doc.add_heading("Common Defects:", 2)
        if documentation.get("common_defects"):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Defect"
            hdr[1].text = "Occurrence"
            for d in documentation["common_defects"]:
                row = table.add_row().cells
                row[0].text = d["defect"]
                row[1].text = d["occurrence"]

        doc.add_heading("Appendix:", 2)
        if documentation.get("appendix_table"):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Parameter"
            hdr[1].text = "Value"
            for p in documentation["appendix_table"]:
                row = table.add_row().cells
                row[0].text = p["parameter"]
                row[1].text = p["value"]

        doc.save(output_path)
        logger.info(f"✅ Word saved to: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Word Error: {e}")
        traceback.print_exc()
        return None

# -------------------------------
# Generate and Save All
# -------------------------------

def generate_and_save_bpg(user_story_json: str) -> Tuple[Optional[str], Optional[str]]:
    pdf_path = None
    word_path = None
    png_path = None
    try:
        user_story = json.loads(user_story_json)
        documentation = generate_documentation(user_story)
        if not documentation:
            raise Exception("Failed to generate documentation content.")

        mermaid_code = documentation.get("flowchart")
        if mermaid_code:
            png_path = render_flowchart_to_files(mermaid_code)

        pdf_path = create_bpg_pdf(documentation, png_path)
        word_path = create_bpg_word(documentation, png_path)

        return pdf_path, word_path
    except Exception as e:
        logger.error(f"Unexpected error in generate_and_save_bpg: {e}")
        traceback.print_exc()
        return None, None

# -------------------------------
# Generate and Save All (MCP Tool)
# -------------------------------

@mcp.tool(description="Generate Business Process Guide (BPG) documents from a user story JSON.")
async def generate_bpg_mcp(user_story: Dict[str, Any]) -> Dict[str, Any]:
    logger.debug(f"Processing user story via MCP: {user_story}")
    pdf_path, word_path = generate_and_save_bpg(json.dumps(user_story))  # Convert dict to JSON string
    return {
        "pdf_path": pdf_path,
        "word_path": word_path,
        "status": "success" if pdf_path and word_path else "failed"
    }

# FastAPI Endpoint
@app.post("/generate_bpg")
async def generate_bpg_endpoint(user_story: UserStoryInput = Body(...)):
    try:
        logger.debug(f"Received request body at endpoint: {user_story.model_dump()}")
        user_story_dict = user_story.model_dump()
        
        # Simulate MCP request
        mcp_request = {
            "tool": "generate_bpg_mcp",
            "params": {"user_story": user_story_dict}
        }
        
        # Process using MCP tool
        response = await generate_bpg_mcp(**mcp_request["params"])
        logger.debug(f"Endpoint response: {response}")
        return response
    except json.JSONDecodeError:
        logger.error("Invalid JSON in request body")
        return {"error": "Invalid JSON in request body", "status": "failed"}
    except Exception as e:
        logger.error(f"Endpoint error: {str(e)}")
        return {"error": str(e), "status": "failed"}

# Function to run the FastAPI server in a separate thread
def run_server():
    config = uvicorn.Config(app, host="127.0.0.1", port=8000, log_level="info")
    server = uvicorn.Server(config)
    server.run()

# Client runner
async def run_client():
    logger.info("Running client...")
    print("Enter your user story details (or paste valid JSON matching {\"Title\": str, \"Description\": str, \"AcceptanceCriteria\": [], \"AdditionalInformation\": str}):")

    try:
        # Attempt to get input as JSON first
        user_input = input().strip()
        if user_input and user_input.startswith("{"):
            user_story = json.loads(user_input)
        else:
            # Fall back to interactive input if not JSON
            title = input("Title: ").strip()
            description = input("Description: ").strip()
            
            print("Enter acceptance criteria (one per line, press Enter twice to finish):")
            acceptance_criteria = []
            while True:
                line = input().strip()
                if line == "":
                    if acceptance_criteria:  # Ensure at least one criterion is entered before finishing
                        break
                    else:
                        print("Please enter at least one acceptance criterion.")
                        continue
                acceptance_criteria.append(line)
            
            additional_information = input("Additional Information (optional, press Enter to skip): ").strip() or None
            
            user_story = {
                "Title": title,
                "Description": description,
                "AcceptanceCriteria": acceptance_criteria,
                "AdditionalInformation": additional_information
            }
        
        logger.debug(f"Constructed user story: {user_story}")
        print("\n📤 Your User Story:")
        print(json.dumps(user_story, indent=2))
        
        try:
            async with httpx.AsyncClient() as client:
                # Wait for server to start with a more reliable check
                await asyncio.sleep(1)  # Reduced to 1 second as a starting point
                for _ in range(5):  # Retry up to 5 times
                    try:
                        url = "http://127.0.0.1:8000/generate_bpg"
                        logger.debug(f"Sending request to: {url}")
                        response = await client.post(url, json=user_story)
                        response.raise_for_status()
                        result = response.json()
                        logger.debug(f"Client received response: {result}")
                        print("✅ BPG Generation Result:")
                        print(json.dumps(result, indent=2))
                        if result.get("pdf_path") and result.get("word_path"):
                            print(f"📄 PDF saved: {result['pdf_path']}")
                            print(f"📝 Word saved: {result['word_path']}")
                        else:
                            print("❌ Generation failed.")
                        break
                    except httpx.RequestError:
                        await asyncio.sleep(1)
                        continue
                else:
                    raise httpx.RequestError("Server did not start in time")
        except json.JSONDecodeError:
            print("❌ Error: Invalid JSON input")
            logger.error(f"Invalid JSON input: {user_story}")
        except httpx.RequestError as e:
            print(f"❌ Error connecting to server: {str(e)}. Is the server running?")
            logger.error(f"Connection error: {str(e)}")
        except Exception as e:
            print(f"❌ Error: {str(e)}")
            logger.error(f"Client error: {str(e)}")

    except json.JSONDecodeError as e:
        print(f"❌ Error: Invalid JSON input - {str(e)}. Please ensure the input is valid JSON.")
        logger.error(f"JSON decode error in client input: {user_input} (Error: {e})")
    except KeyboardInterrupt:
        print("\n❌ Process interrupted by user.")
        logger.info("Client process interrupted.")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        logger.error(f"Unexpected client error: {e}")

# Main function to run server and client
async def main():
    # Start FastAPI server in a separate thread
    server_thread = threading.Thread(target=run_server, daemon=True)
    logger.info("Starting FastAPI server on http://127.0.0.1:8000 in a separate thread")
    server_thread.start()
    
    # Run client
    await run_client()

# CLI Entry Point (Optional, for file-based input)
def run_cli():
    parser = argparse.ArgumentParser(description="Generate BPG documents from user story JSON.")
    parser.add_argument("json_file", help="Path to user story JSON file")
    args = parser.parse_args()

    try:
        with open(args.json_file, "r") as f:
            user_story_json = f.read()
        print("⚙️ Generating Business Process Guide...")
        pdf_path, word_path = generate_and_save_bpg(user_story_json)
        if pdf_path and word_path:
            print(f"✅ PDF: {pdf_path}")
            print(f"✅ Word: {word_path}")
        else:
            print("❌ Failed to generate one or both documents.")
    except Exception as e:
        print(f"❌ Error: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    # Ensure Windows compatibility
    if os.name == "nt":
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    
    # Check for CLI arguments
    if len(sys.argv) > 1:
        run_cli()
    else:
        # Run server and client by default
        asyncio.run(main())