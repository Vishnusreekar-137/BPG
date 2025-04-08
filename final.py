import os
import json
import requests
from langchain_anthropic import ChatAnthropic
from langchain.prompts import PromptTemplate
from langchain_core.runnables import RunnableSequence
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
import traceback
from docx.shared import Inches

# Load environment variables
load_dotenv()

# --- Step 1: Generate Documentation ---

def generate_documentation(user_story_json):
    """Generates documentation content from JSON input."""
    try:
        # Parse and validate JSON
        try:
            user_story_data = json.loads(user_story_json)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format in file: {e}. Check commas, brackets, or quotes.")

        # Extract description from nested structure
        if "UserStory" in user_story_data and "Description" in user_story_data["UserStory"]:
            description = user_story_data["UserStory"]["Description"]
        elif "description" in user_story_data:
            description = user_story_data["description"]
        else:
            raise ValueError("JSON must contain a 'description' key or a 'UserStory' object with a 'Description' key.")

        # Generate Mermaid syntax
        mermaid_syntax = generate_mermaid(description)
        if not mermaid_syntax:
            print("Warning: Failed to generate Mermaid syntax. Flowchart will be missing.")

        # Sample data for Common Defects and Appendix Table
        common_defects = [
            {"defect": "Invalid email format", "occurrence": "15%"},
            {"defect": "Password too short", "occurrence": "10%"},
            {"defect": "Email already registered", "occurrence": "5%"}
        ]
        appendix_table = [
            {"parameter": "Email Validation Regex", "value": r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$"},
            {"parameter": "Password Min Length", "value": "8"},
            {"parameter": "Confirmation Email Template", "value": "welcome_email.html"}
        ]

        return {
            "description": description,
            "flowchart": mermaid_syntax,
            "common_defects": common_defects,
            "appendix_table": appendix_table
        }
    except Exception as e:
        print(f"Error in generate_documentation: {e}")
        traceback.print_exc()
        return None

def generate_mermaid(description, theme="forest"):
    """Generates Mermaid syntax using LangChain with Claude 3 Haiku."""
    api_key = os.environ.get("CLAUDE_API_KEY")
    if not api_key:
        print("Error: CLAUDE_API_KEY not found.")
        return None

    try:
        llm = ChatAnthropic(
            model_name="claude-3-haiku-20240307",
            anthropic_api_key=api_key,
            max_tokens=1000,
            temperature=0.1
        )

        prompt_template = PromptTemplate(
            input_variables=["description"],
            template="Generate only the Mermaid syntax for a flowchart representing this process: {description}. Return ONLY the syntax starting with 'graph', with no additional text, comments, or formatting. Use 'finish' instead of 'end' for the final node. For example, for a process 'A leads to B', return: graph TD; A-->B; B-->finish;"
        )

        chain = RunnableSequence(prompt_template | llm)
        raw_content = chain.invoke({"description": description}).content.strip()

        # Clean and validate the Mermaid code
        def clean_mermaid_code(content):
            lines = content.split('\n')
            valid_lines = []
            for i, line in enumerate(lines):
                line = line.strip()
                if line and (line.startswith("graph") or line.startswith("flowchart") or valid_lines):
                    # Replace 'end' with 'finish' in node names
                    if 'end' in line.lower() and ('[' in line or '{' in line or '-->' in line):
                        line = line.replace('end', 'finish').replace('End', 'Finish')
                    valid_lines.append(line)
            cleaned = '\n'.join(valid_lines)
            return cleaned if cleaned.startswith(("graph", "flowchart")) else None

        mermaid_code = clean_mermaid_code(raw_content)
        if not mermaid_code:
            print(f"Warning: Generated content is not valid Mermaid syntax: {raw_content[:150]}...")
            return None

        init_block = f"%%{{init: {{'theme': '{theme}'}}}}%%"
        if theme and mermaid_code.startswith(("graph", "flowchart")) and "%%{init:" not in mermaid_code:
            return f"{init_block}\n{mermaid_code}"
        return mermaid_code
    except Exception as e:
        print(f"Error generating Mermaid syntax: {e}")
        traceback.print_exc()
        return None

# --- Step 3: Create PDF with Improved Layout ---

def create_bpg_pdf(documentation, image_path, output_path="business_process_guide.pdf"):
    """Creates a PDF with description and flowchart on page 1, tables on page 2."""
    try:
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # Unified spacing configuration
        section_spacing = 0.1 * inch
        table_spacing = 0.05 * inch

        # Title Section (Page 1)
        elements.append(Paragraph("Business Process Guide", styles['Heading1']))
        elements.append(Spacer(1, 0.15 * inch))

        # Description Section (Page 1)
        elements.append(Paragraph("Description:", styles['Heading2']))
        desc_style = styles['Normal']
        desc_style.fontSize = 10  # Reduce font size to fit with flowchart
        desc_style.leading = 12   # Adjust line spacing
        elements.append(Paragraph(documentation.get("description", "No description available"), desc_style))
        elements.append(Spacer(1, section_spacing))

        # Process Flow Section (Page 1)
        elements.append(Paragraph("Process Flow:", styles['Heading2']))
        if image_path and os.path.exists(image_path):
            try:
                # Adjust flowchart size to fit on page 1 with description
                img = Image(image_path, width=4*inch, height=2.5*inch, kind='proportional')
                elements.append(img)
            except Exception as e:
                elements.append(Paragraph("[Image Error]", styles['Italic']))
        elements.append(Spacer(1, section_spacing))

        # Force a page break after flowchart to move tables to page 2
        elements.append(PageBreak())  # Corrected from doc.PageBreak()

        # Common Defects Table (Page 2)
        elements.append(Paragraph("Common Defects:", styles['Heading2']))
        if documentation.get("common_defects"):
            defect_data = [["Defect", "Occurrence"]] + [
                [d["defect"], d["occurrence"]] for d in documentation["common_defects"]
            ]
            defect_table = Table(defect_data, colWidths=[3.5*inch, 1.5*inch])
            defect_table.setStyle(TableStyle([
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 9),
                ('FONTSIZE', (0,1), (-1,-1), 8),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LINEBELOW', (0,0), (-1,0), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.whitesmoke, colors.white]),
                ('BOX', (0,0), (-1,-1), 0.5, colors.black),
                ('LEFTPADDING', (0,0), (-1,-1), 3),
                ('RIGHTPADDING', (0,0), (-1,-1), 3),
            ]))
            elements.append(defect_table)
            elements.append(Spacer(1, table_spacing))
        
        # Appendix Table (Page 2)
        elements.append(Paragraph("Appendix:", styles['Heading2']))
        if documentation.get("appendix_table"):
            appendix_data = [["Parameter", "Value"]] + [
                [p["parameter"], p["value"]] for p in documentation["appendix_table"]
            ]
            appendix_table = Table(appendix_data, colWidths=[3*inch, 3*inch])
            appendix_table.setStyle(TableStyle([
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 9),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LINEBELOW', (0,0), (-1,0), 0.5, colors.grey),
                ('BOX', (0,0), (-1,-1), 0.5, colors.lightgrey),
                ('LEFTPADDING', (0,0), (-1,-1), 3),
                ('GRID', (0,0), (-1,-1), 0.5, colors.lightgrey),
            ]))
            elements.append(appendix_table)
            elements.append(Spacer(1, table_spacing))

        doc.build(elements)
        return output_path
    except Exception as e:
        print(f"PDF Error: {e}")
        return None

# --- Step 4: Create Word Document with Improved Layout ---

def create_bpg_word(documentation, png_path, output_path="business_process_guide.docx"):
    """Creates a Word doc with description and flowchart on page 1, tables on page 2."""
    try:
        # Check if the output file is accessible
        if os.path.exists(output_path):
            try:
                with open(output_path, 'rb') as f:
                    pass  # Test if file can be opened
            except PermissionError:
                raise PermissionError(f"Cannot save to '{output_path}': File is locked or open in another application. Please close it and retry.")

        doc = Document()
        
        # Set document-wide styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)

        # Title Section (Page 1)
        title = doc.add_heading("Business Process Guide", 1)
        title.paragraph_format.space_after = Pt(6)

        # Description Section (Page 1)
        doc.add_heading("Description:", 2).paragraph_format.space_after = Pt(3)
        desc_para = doc.add_paragraph(documentation.get("description", "No description available"))
        desc_para.runs[0].font.size = Pt(9)  # Reduce font size to fit with flowchart

        # Process Flow Section (Page 1)
        flow_heading = doc.add_heading("Process Flow:", 2)
        flow_heading.paragraph_format.space_after = Pt(3)
        if png_path and os.path.exists(png_path):
            # Adjust flowchart size to fit on page 1 with description
            doc.add_picture(png_path, width=Inches(4), height=Inches(2.5))

        # Add page break to move tables to page 2
        doc.add_page_break()

        # Common Defects Table (Page 2)
        doc.add_heading("Common Defects:", 2).paragraph_format.space_after = Pt(3)
        if documentation.get("common_defects"):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.autofit = False
            
            # Set column widths
            table.columns[0].width = Inches(3.5)
            table.columns[1].width = Inches(1.5)
            
            # Header row
            hdr = table.rows[0].cells
            hdr[0].text = "Defect"
            hdr[1].text = "Occurrence"
            for cell in hdr:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for defect in documentation["common_defects"]:
                row = table.add_row().cells
                row[0].text = defect["defect"]
                row[1].text = defect["occurrence"]
                for cell in row:
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        # Appendix Table (Page 2)
        doc.add_heading("Appendix:", 2).paragraph_format.space_after = Pt(3)
        if documentation.get("appendix_table"):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.autofit = False
            
            # Set column widths
            table.columns[0].width = Inches(3)
            table.columns[1].width = Inches(3)
            
            # Header row
            hdr = table.rows[0].cells
            hdr[0].text = "Parameter"
            hdr[1].text = "Value"
            for cell in hdr:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for param in documentation["appendix_table"]:
                row = table.add_row().cells
                row[0].text = param["parameter"]
                row[1].text = param["value"]
                for cell in row:
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        doc.save(output_path)
        return output_path
    except PermissionError as e:
        print(f"Word Error: {e}")
        return None
    except Exception as e:
        print(f"Word Error: {e}")
        return None

# --- Main Function ---

# --- Step 2: Render Flowchart to PNG ---
def render_flowchart_to_files(mermaid_code, png_file="flowchart.png", timeout=30):
    """Renders Mermaid code to PNG using Kroki."""
    if not mermaid_code:
        print("No Mermaid code provided to render.")
        return None

    # Clean the Mermaid code by removing trailing invalid lines (e.g., '---')
    cleaned_code = '\n'.join(line for line in mermaid_code.split('\n') if line.strip() and not line.strip().startswith('---'))
    cleaned_code = cleaned_code.strip()

    try:
        print(f"Sending Mermaid code to Kroki:\n---\n{cleaned_code}\n---")
        response = requests.post(
            "https://kroki.io/mermaid/png",
            data=cleaned_code.encode("utf-8"),
            headers={"Content-Type": "text/plain; charset=utf-8"},
            timeout=timeout
        )
        response.raise_for_status()

        with open(png_file, "wb") as f:
            f.write(response.content)
        print(f"Flowchart PNG saved successfully to {png_file}")
        return png_file
    except requests.exceptions.RequestException as e:
        print(f"Error rendering PNG via Kroki: {e}")
        if response.text:
            print(f"Kroki error message: {response.text}")
        traceback.print_exc()
        return None
    except Exception as e:
        print(f"Unexpected error in render_flowchart_to_files: {e}")
        traceback.print_exc()
        return None

# --- Main Function ---
def generate_and_save_bpg(user_story_json):
    """Orchestrates the generation of documents."""
    pdf_path = None
    word_path = None
    png_path = None

    try:
        print("Step 1: Generating documentation...")
        documentation = generate_documentation(user_story_json)
        if not documentation:
            raise Exception("Failed to generate documentation content.")

        print("Step 2: Rendering flowchart...")
        mermaid_code = documentation.get("flowchart")
        if mermaid_code:
            png_path = render_flowchart_to_files(mermaid_code)  # Call the defined function
            if not png_path:
                print("Warning: Flowchart rendering failed. Proceeding without image.")
        else:
            print("No flowchart generated.")

        print("Step 3: Creating PDF...")
        pdf_path = create_bpg_pdf(documentation, png_path)

        print("Step 4: Creating Word document...")
        word_path = create_bpg_word(documentation, png_path)

        return pdf_path, word_path
    except Exception as e:
        print(f"Unexpected error in generate_and_save_bpg: {e}")
        traceback.print_exc()
        return None, None

# --- FastAPI Integration ---

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import base64
import uvicorn

app = FastAPI()

class UserStoryInput(BaseModel):
    data: dict

@app.post("/generate-bpg")
async def generate_bpg(input_data: UserStoryInput):
    """
    API endpoint to generate PDF and Word documents from user story JSON and return them encoded in base64.
    """
    try:
        # Convert the input dict to a JSON string as expected by generate_documentation
        user_story_json = json.dumps(input_data.data)

        # Generate the documents using the existing function
        pdf_path, word_path = generate_and_save_bpg(user_story_json)

        # Check if generation was successful
        if not pdf_path or not word_path:
            raise HTTPException(status_code=500, detail="Failed to generate one or both documents.")

        # Read and encode the PDF
        with open(pdf_path, "rb") as pdf_file:
            pdf_bytes = pdf_file.read()
            pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')

        # Read and encode the Word document
        with open(word_path, "rb") as word_file:
            word_bytes = word_file.read()
            word_base64 = base64.b64encode(word_bytes).decode('utf-8')

        # Return the encoded documents in JSON
        return {
            "pdf": pdf_base64,
            "word": word_base64
        }
    except Exception as e:
        print(f"API Error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

# --- Main Execution ---

if __name__ == "__main__":
    if not os.environ.get("CLAUDE_API_KEY"):
        print("Error: CLAUDE_API_KEY not found in .env file.")
        exit(1)
    
    print("Starting FastAPI server for Business Process Guide generation...")
    uvicorn.run(app, host="127.0.0.1", port=8000)