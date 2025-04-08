import os
import json
import requests
import traceback
from dotenv import load_dotenv
from langchain_anthropic import ChatAnthropic
from langchain_together import ChatTogether
from langchain.prompts import PromptTemplate
from langchain_core.runnables import RunnableSequence
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches

# Load environment variables
load_dotenv()

# Initialize LLMs
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY")
TOGETHER_API_KEY = os.getenv("TOGETHER_API_KEY")

claude_llm = ChatAnthropic(
    model_name="claude-3-haiku-20240307",
    anthropic_api_key=CLAUDE_API_KEY,
    temperature=0.1,
    max_tokens=1000
)

together_llm = ChatTogether(
    model="mistralai/Mixtral-8x7B-Instruct-v0.1",
    together_api_key=TOGETHER_API_KEY,
    temperature=0.2
)

# -------------------------------
# Step 1: Generate Documentation
# -------------------------------

def generate_documentation(user_story_json):
    try:
        user_story_data = json.loads(user_story_json)

        if "UserStory" in user_story_data and "Description" in user_story_data["UserStory"]:
            description = user_story_data["UserStory"]["Description"]
        elif "description" in user_story_data:
            description = user_story_data["description"]
        else:
            raise ValueError("JSON must contain a 'description' key or a 'UserStory' object with a 'Description' key.")

        # 🧠 Step 1: Generate Mermaid Flowchart (Claude)
        mermaid_syntax = generate_mermaid(description)

        # 🧠 Step 2: Use Together to infer common defects
        defect_prompt = PromptTemplate(
            input_variables=["description"],
            template="""From the following feature description, identify potential edge cases or user-related defects that may occur during implementation or usage. Return a JSON array like:
[
  {{ "defect": "...", "occurrence": "..." }},
  ...
]
Description:
{description}"""
        )
        defect_chain = RunnableSequence(defect_prompt | together_llm)
        defects_json = defect_chain.invoke({"description": description}).content.strip()
        common_defects = json.loads(defects_json)

        # 🧠 Step 3: Use Together to infer Appendix table data
        appendix_prompt = PromptTemplate(
            input_variables=["description"],
            template="""Based on the feature below, extract relevant configuration parameters or rules that developers or testers should be aware of. Return JSON like:
[
  {{ "parameter": "...", "value": "..." }},
  ...
]
Description:
{description}"""
        )
        appendix_chain = RunnableSequence(appendix_prompt | together_llm)
        appendix_json = appendix_chain.invoke({"description": description}).content.strip()
        appendix_table = json.loads(appendix_json)

        return {
            "description": description,
            "flowchart": mermaid_syntax,
            "common_defects": common_defects,
            "appendix_table": appendix_table
        }

    except Exception as e:
        print(f"Error in generate_documentation: {e}")
        traceback.print_exc()
        return None
# -------------------------------
# Step 1a: Generate Mermaid with Claude
# -------------------------------

def generate_mermaid(description, theme="forest"):
    try:
        prompt_template = PromptTemplate(
            input_variables=["description"],
            template="Generate only the Mermaid syntax for a flowchart representing this process: {description}. Return ONLY the syntax starting with 'graph', with no additional text, comments, or formatting. Use 'finish' instead of 'end'."
        )

        chain = RunnableSequence(prompt_template | claude_llm)
        raw_content = chain.invoke({"description": description}).content.strip()

        def clean_mermaid_code(content):
            lines = content.split('\n')
            valid_lines = []
            for line in lines:
                line = line.strip()
                if line and (line.startswith("graph") or line.startswith("flowchart") or valid_lines):
                    if 'end' in line.lower():
                        line = line.replace('end', 'finish').replace('End', 'Finish')
                    valid_lines.append(line)
            cleaned = '\n'.join(valid_lines)
            return cleaned if cleaned.startswith(("graph", "flowchart")) else None

        mermaid_code = clean_mermaid_code(raw_content)
        if not mermaid_code:
            print("Warning: Generated content is not valid Mermaid syntax.")
            return None

        init_block = f"%%{{init: {{'theme': '{theme}'}}}}%%"
        if "%%{init:" not in mermaid_code:
            return f"{init_block}\n{mermaid_code}"
        return mermaid_code
    except Exception as e:
        print(f"Error generating Mermaid syntax: {e}")
        traceback.print_exc()
        return None

# -------------------------------
# Step 2: Render Mermaid PNG via Kroki
# -------------------------------

def render_flowchart_to_files(mermaid_code, png_file="flowchart.png", timeout=30):
    if not mermaid_code:
        print("No Mermaid code provided.")
        return None

    cleaned_code = '\n'.join(line for line in mermaid_code.split('\n') if line.strip() and not line.strip().startswith('---'))

    try:
        response = requests.post(
            "https://kroki.io/mermaid/png",
            data=cleaned_code.encode("utf-8"),
            headers={"Content-Type": "text/plain; charset=utf-8"},
            timeout=timeout
        )
        response.raise_for_status()

        with open(png_file, "wb") as f:
            f.write(response.content)
        print(f"✅ Flowchart PNG saved to: {png_file}")
        return png_file
    except Exception as e:
        print(f"Error rendering PNG: {e}")
        traceback.print_exc()
        return None

# -------------------------------
# Step 3: Create PDF
# -------------------------------

def create_bpg_pdf(documentation, image_path, output_path="business_process_guide.pdf"):
    try:
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        normal = styles['Normal']
        heading = styles['Heading2']
        elements = []

        # Title
        elements.append(Paragraph("Business Process Guide", styles['Heading1']))
        elements.append(Spacer(1, 0.2 * inch))

        # Description
        elements.append(Paragraph("Description:", heading))
        elements.append(Paragraph(documentation.get("description", "No description available"), normal))
        elements.append(Spacer(1, 0.15 * inch))

        # Process Flow
        elements.append(Paragraph("Process Flow:", heading))
        if image_path and os.path.exists(image_path):
            img = Image(image_path, width=4.5*inch, height=2.7*inch)
            elements.append(img)
        elements.append(PageBreak())

        # Common Defects Table
        elements.append(Paragraph("Common Defects:", heading))
        if documentation.get("common_defects"):
            defect_data = [["Defect", "Occurrence"]]
            for d in documentation["common_defects"]:
                # Use Paragraph for both columns to enable text wrapping
                defect_data.append([Paragraph(d.get("defect", ""), normal), 
                                    Paragraph(d.get("occurrence", ""), normal)])
            
            # Adjusted column widths for better space distribution
            defect_table = Table(defect_data, colWidths=[2.0*inch, 4.4*inch])
            
            # Table styling with added padding for readability
            defect_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Bold header
            ('FONTSIZE', (0, 0), (-1, -1), 9),                # Font size 9 for the table
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),              # Left alignment
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),              # Top vertical alignment
            ('LEFTPADDING', (0, 0), (-1, -1), 4),             # Left padding for cells
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),            # Right padding for cells
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),           # Bottom padding for cells
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),    # Grey grid lines
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),     # Ensure header text is visible
    ]))
            elements.append(defect_table)
            elements.append(Spacer(1, 0.2 * inch))

        # Appendix Table
        elements.append(Paragraph("Appendix:", heading))
        if documentation.get("appendix_table"):
            appendix_data = [["Parameter", "Value"]]
            for p in documentation["appendix_table"]:
                appendix_data.append([
                    Paragraph(p.get("parameter", ""), normal),
                    Paragraph(p.get("value", ""), normal)
                ])
            appendix_table = Table(appendix_data, colWidths=[2.2*inch, 4.2*inch])
            appendix_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            elements.append(appendix_table)

        doc.build(elements)
        return output_path

    except Exception as e:
        print(f"PDF Error: {e}")
        traceback.print_exc()
        return None
# -------------------------------
# Step 4: Create Word
# -------------------------------

def create_bpg_word(documentation, png_path, output_path="business_process_guide.docx"):
    try:
        if os.path.exists(output_path):
            with open(output_path, 'rb'):
                pass

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        doc.add_heading("Business Process Guide", 1)
        doc.add_heading("Description:", 2)
        doc.add_paragraph(documentation.get("description", "No description available"))

        doc.add_heading("Process Flow:", 2)
        if png_path and os.path.exists(png_path):
            doc.add_picture(png_path, width=Inches(6), height=Inches(6.5))

        doc.add_page_break()

        doc.add_heading("Common Defects:", 2)
        if documentation.get("common_defects"):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Defect"
            hdr[1].text = "Occurrence"
            for d in documentation["common_defects"]:
                row = table.add_row().cells
                row[0].text = d["defect"]
                row[1].text = d["occurrence"]

        doc.add_heading("Appendix:", 2)
        if documentation.get("appendix_table"):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Parameter"
            hdr[1].text = "Value"
            for p in documentation["appendix_table"]:
                row = table.add_row().cells
                row[0].text = p["parameter"]
                row[1].text = p["value"]

        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Word Error: {e}")
        traceback.print_exc()
        return None

# -------------------------------
# Generate and Save All
# -------------------------------

def generate_and_save_bpg(user_story_json):
    pdf_path = None
    word_path = None
    png_path = None
    try:
        documentation = generate_documentation(user_story_json)
        if not documentation:
            raise Exception("Failed to generate documentation content.")

        mermaid_code = documentation.get("flowchart")
        if mermaid_code:
            png_path = render_flowchart_to_files(mermaid_code)

        pdf_path = create_bpg_pdf(documentation, png_path)
        word_path = create_bpg_word(documentation, png_path)

        return pdf_path, word_path
    except Exception as e:
        print(f"Unexpected error in generate_and_save_bpg: {e}")
        traceback.print_exc()
        return None, None

# -------------------------------
# CLI Entry Point
# -------------------------------

if __name__ == "__main__":
    import argparse

    if not CLAUDE_API_KEY or not TOGETHER_API_KEY:
        print("❌ Please set CLAUDE_API_KEY and TOGETHER_API_KEY in your .env file.")
        exit(1)

    parser = argparse.ArgumentParser(description="Generate BPG documents from user story JSON.")
    parser.add_argument("json_file", help="Path to user story JSON file")
    args = parser.parse_args()

    try:
        with open(args.json_file, "r") as f:
            user_story_json = f.read()
        print("⚙️ Generating Business Process Guide...")
        pdf_path, word_path = generate_and_save_bpg(user_story_json)
        if pdf_path and word_path:
            print(f"✅ PDF: {pdf_path}")
            print(f"✅ Word: {word_path}")
        else:
            print("❌ Failed to generate one or both documents.")
    except Exception as e:
        print(f"❌ Error: {e}")
        traceback.print_exc()
