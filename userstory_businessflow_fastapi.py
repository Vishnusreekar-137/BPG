from together import Together
import os
from dotenv import load_dotenv
import uuid
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import json
import sys
import subprocess
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from cairosvg import svg2png
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.units import inch
from mcp.server.fastmcp import FastMCP
import logging
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import Optional, List
import uvicorn

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("BPG-Analyzer")

# Load environment variables from .env file
load_dotenv()

# Get API key from environment variables
api_key = os.getenv("TOGETHER_API_KEY")
if not api_key:
    raise ValueError("No TOGETHER_API_KEY found in environment variables")

# Initialize Together client
client = Together(api_key=api_key)

# Initialize MCP server
mcp = FastMCP("bpg-generator")

# Initialize FastAPI app
app = FastAPI(title="BPG Generator API", description="Generate Business Process Guides from user stories.")

# Pydantic model for FastAPI request
class UserStoryRequest(BaseModel):
    Title: str
    Description: str
    AcceptanceCriteria: List[str]
    AdditionalInformation: str

# Pydantic model for FastAPI response
class BPGResponse(BaseModel):
    word_path: Optional[str]
    pdf_path: Optional[str]
    html_path: Optional[str]
    json_path: Optional[str]
    status: str

def generate_documentation(user_story_data: dict):
    """Generate documentation including LLM-analyzed description, common defects, and appendix table"""
    try:
        # Use Description as the primary user story
        user_story = user_story_data.get("Description", "")
        if not user_story:
            logger.error("No Description provided in user story data")
            return None

        # Enhance with AcceptanceCriteria and AdditionalInformation
        acceptance_criteria = user_story_data.get("AcceptanceCriteria", [])
        additional_info = user_story_data.get("AdditionalInformation", "")
        full_context = f"{user_story}\n\nAcceptance Criteria:\n" + "\n".join(f"- {crit}" for crit in acceptance_criteria)
        if additional_info:
            full_context += f"\n\nAdditional Information: {additional_info}"

        # Generate an analyzed description
        description_prompt = f"""
        Analyze the following user story, including its acceptance criteria and additional information, and provide a concise, polished description summarizing its purpose and key details. Do not return the raw input text, but instead a refined narrative based on it:

        User Story:
        {full_context}

        Return the description as plain text with no additional formatting or labels.
        """
        response = client.chat.completions.create(
            model="meta-llama/Llama-4-Scout-17B-16E-Instruct",
            temperature=0.2,
            messages=[{"role": "user", "content": description_prompt}]
        )
        description = response.choices[0].message.content.strip()
        logger.debug(f"Generated description: {description}")

        # Infer common defects
        defect_prompt = f"""
        From the following user story, including its acceptance criteria and additional information, identify potential edge cases or user-related defects that may occur during implementation or usage. Return **only** a valid JSON array with no additional text, markdown, or explanations. The JSON must follow this exact format:
        [
          {{"defect": "string", "occurrence": "string"}},
          ...
        ]
        If no defects are identified, return an empty array [].
        User Story:
        {full_context}
        """
        response = client.chat.completions.create(
            model="meta-llama/Llama-4-Scout-17B-16E-Instruct",
            temperature=0.2,
            messages=[{"role": "user", "content": defect_prompt}]
        )
        defects_content = response.choices[0].message.content.strip()
        logger.debug(f"Raw defects response: {defects_content}")

        # Try direct JSON parsing first
        common_defects = []
        try:
            common_defects = json.loads(defects_content)
            if not isinstance(common_defects, list):
                logger.warning(f"Defects response is not a JSON array: {defects_content}")
                common_defects = []
        except json.JSONDecodeError:
            # Fallback to regex
            json_matches = re.findall(r'\[\s*\{.*?\}\s*\]', defects_content, re.DOTALL)
            if json_matches:
                try:
                    common_defects = json.loads(json_matches[0])
                    if not isinstance(common_defects, list):
                        common_defects = []
                except json.JSONDecodeError as e:
                    logger.error(f"Failed to parse JSON from regex match: {e}")
            else:
                logger.warning(f"No valid JSON array found in defects response: {defects_content}")

        # Validate defect structure
        if not all(isinstance(item, dict) and "defect" in item and "occurrence" in item for item in common_defects):
            logger.warning(f"Invalid defect structure: {common_defects}")
            common_defects = []

        # Infer appendix table data
        appendix_prompt = f"""
        Based on the user story below, including its acceptance criteria and additional information, extract relevant configuration parameters or rules that developers or testers should be aware of. Return **only** a valid JSON array with no additional text, markdown, or explanations. The JSON must follow this exact format:
        [
          {{"parameter": "string", "value": "string"}},
          ...
        ]
        If no parameters are identified, return an empty array [].
        User Story:
        {full_context}
        """
        response = client.chat.completions.create(
            model="meta-llama/Llama-4-Scout-17B-16E-Instruct",
            temperature=0.2,
            messages=[{"role": "user", "content": appendix_prompt}]
        )
        appendix_content = response.choices[0].message.content.strip()
        logger.debug(f"Raw appendix response: {appendix_content}")

        # Try direct JSON parsing first
        appendix_table = []
        try:
            appendix_table = json.loads(appendix_content)
            if not isinstance(appendix_table, list):
                logger.warning(f"Appendix response is not a JSON array: {appendix_content}")
                appendix_table = []
        except json.JSONDecodeError:
            # Fallback to regex
            json_matches = re.findall(r'\[\s*\{.*?\}\s*\]', appendix_content, re.DOTALL)
            if json_matches:
                try:
                    appendix_table = json.loads(json_matches[0])
                    if not isinstance(appendix_table, list):
                        appendix_table = []
                except json.JSONDecodeError as e:
                    logger.error(f"Failed to parse JSON from regex match: {e}")
            else:
                logger.warning(f"No valid JSON array found in appendix response: {appendix_content}")

        # Validate appendix structure
        if not all(isinstance(item, dict) and "parameter" in item and "value" in item for item in appendix_table):
            logger.warning(f"Invalid appendix structure: {appendix_table}")
            appendix_table = []

        return {
            "title": user_story_data.get("Title", ""),
            "description": description,
            "common_defects": common_defects,
            "appendix_table": appendix_table,
            "acceptance_criteria": acceptance_criteria,
            "additional_info": additional_info
        }
    except Exception as e:
        logger.error(f"Error in generate_documentation: {e}")
        return None

def analyze_flow_structure(user_story_data: dict):
    """Extract flow structure directly from the user story for flowchart"""
    user_story = user_story_data.get("Description", "")
    if not user_story:
        logger.error("No Description provided in user story data")
        return []

    # Enhance with AcceptanceCriteria and AdditionalInformation
    acceptance_criteria = user_story_data.get("AcceptanceCriteria", [])
    additional_info = user_story_data.get("AdditionalInformation", "")
    full_context = f"{user_story}\n\nAcceptance Criteria:\n" + "\n".join(f"- {crit}" for crit in acceptance_criteria)
    if additional_info:
        full_context += f"\n\nAdditional Information: {additional_info}"

    prompt = f"""
    Analyze the following User Story, including its acceptance criteria and additional information, to create a detailed flow chart structure:

    USER STORY:
    {full_context}

    Extract the main steps, decision points, and alternative flows from this user story.
    Format your response as **only** a valid JSON array with no additional text, markdown, or explanations. Each node must have:
    - "id": A unique identifier (use snake_case)
    - "text": The text to display in the node
    - "type": Either "process" (for actions/steps), "decision" (for yes/no questions), "start", or "end"
    - "next": Either the ID of the next node, or for decision nodes, an object with "yes" and "no" keys pointing to node IDs

    Start with a "start" node and end with an "end" node.
    Return ONLY the JSON array.
    """
    logger.debug("Analyzing flow structure... This may take a minute.")
    response = client.chat.completions.create(
        model="meta-llama/Llama-4-Scout-17B-16E-Instruct",
        temperature=0.2,
        messages=[{"role": "user", "content": prompt}]
    )
    flow_content = response.choices[0].message.content.strip()
    logger.debug(f"Raw flow response: {flow_content}")

    # Try direct JSON parsing
    try:
        flow_definition = json.loads(flow_content)
        if not isinstance(flow_definition, list):
            logger.error(f"Flow response is not a JSON array: {flow_content}")
            return []
    except json.JSONDecodeError:
        # Fallback to regex
        json_match = re.search(r'\[\s*\{.*\}\s*\]', flow_content, re.DOTALL)
        if json_match:
            try:
                flow_definition = json.loads(json_match.group(0))
                if not isinstance(flow_definition, list):
                    logger.error(f"Parsed flow is not a JSON array: {flow_definition}")
                    return []
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse JSON from regex match: {e}")
                return []
        else:
            # Clean markdown if present
            json_str = re.sub(r'^```(?:json)?\s*|\s*```$', '', flow_content.strip())
            try:
                flow_definition = json.loads(json_str)
                if not isinstance(flow_definition, list):
                    logger.error(f"Cleaned flow is not a JSON array: {flow_definition}")
                    return []
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse cleaned JSON: {e}")
                return []

    # Validate flow structure
    required_keys = {"id", "text", "type"}
    for node in flow_definition:
        if not all(key in node for key in required_keys):
            logger.warning(f"Invalid node structure: {node}")
            return []
        if node["type"] not in {"start", "end", "process", "decision"}:
            logger.warning(f"Invalid node type in: {node}")
            return []
        if node["type"] == "decision" and not isinstance(node.get("next"), dict):
            logger.warning(f"Decision node missing yes/no next: {node}")
            return []

    return flow_definition

def fix_missing_connections(flow_definition):
    """Fix missing connections in the flow definition"""
    node_map = {node['id']: node for node in flow_definition}
    start_nodes = [node for node in flow_definition if node.get('type') == 'start']
    if start_nodes and 'next' not in start_nodes[0]:
        for node in flow_definition:
            if node.get('type') != 'start' and node.get('type') != 'end':
                start_nodes[0]['next'] = node['id']
                break
    for node in flow_definition:
        if node.get('type') != 'end' and 'next' not in node:
            for potential_next in flow_definition:
                if potential_next['id'] != node['id'] and potential_next.get('type') != 'start':
                    node['next'] = potential_next['id']
                    break

def json_to_pyflowchart(flow_definition, output_dir, generation_id):
    """Convert JSON flow definition to PyFlowchart code"""
    fix_missing_connections(flow_definition)
    node_map = {node['id']: node for node in flow_definition}

    py_script = [
        "from pyflowchart import *",
        "",
        "# Create nodes",
    ]

    for node in flow_definition:
        node_id = node['id']
        node_text = node['text'].replace('"', '\\"')
        node_type = node.get('type', 'process')
        if node_type == 'start':
            py_script.append(f'{node_id} = StartNode("{node_text}")')
        elif node_type == 'end':
            py_script.append(f'{node_id} = EndNode("{node_text}")')
        elif node_type == 'decision':
            py_script.append(f'{node_id} = ConditionNode("{node_text}")')
        else:
            py_script.append(f'{node_id} = OperationNode("{node_text}")')

    py_script.append("")
    py_script.append("# Create connections")

    for node in flow_definition:
        node_id = node['id']
        if isinstance(node.get('next'), str) and node.get('next'):
            next_id = node['next']
            if next_id in node_map:
                py_script.append(f'{node_id}.connect({next_id})')
        elif isinstance(node.get('next'), dict):
            yes_id = node.get('next', {}).get('yes')
            no_id = node.get('next', {}).get('no')
            if yes_id and yes_id in node_map:
                py_script.append(f'{node_id}.connect_yes({yes_id})')
            if no_id and no_id in node_map:
                py_script.append(f'{node_id}.connect_no({no_id})')

    start_nodes = [node for node in flow_definition if node.get('type') == 'start']
    start_node_id = start_nodes[0]['id'] if start_nodes else flow_definition[0]['id']

    py_script.append("")
    py_script.append("# Create flowchart")
    py_script.append(f"fc = Flowchart({start_node_id})")
    py_script.append('print(fc.flowchart())')

    script_path = os.path.join(output_dir, f"flowchart_script_{generation_id}.py")
    with open(script_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(py_script))

    logger.debug(f"PyFlowchart script saved to: {script_path}")

    try:
        import pyflowchart
        logger.debug("Using installed PyFlowchart package")
    except ImportError:
        logger.debug("PyFlowchart not installed. Installing...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pyflowchart"])
        logger.debug("PyFlowchart installed successfully")

    try:
        result = subprocess.run(
            [sys.executable, script_path],
            capture_output=True,
            text=True,
            check=True
        )
        flowchart_code = result.stdout.strip()
    except subprocess.CalledProcessError as e:
        logger.error(f"Failed to run PyFlowchart script: {e}")
        return None, script_path, None

    flowchart_path = os.path.join(output_dir, f"flowchart_{generation_id}.txt")
    with open(flowchart_path, 'w', encoding='utf-8') as f:
        f.write(flowchart_code)

    logger.debug(f"Flowchart code saved to: {flowchart_path}")
    return flowchart_code, script_path, flowchart_path

def generate_flowchart_image(flowchart_code, output_dir, generation_id):
    """Generate an SVG flowchart and convert it to high-resolution PNG"""
    if not flowchart_code:
        logger.error("No flowchart code provided")
        return None, None

    html_path = os.path.join(output_dir, f"flowchart_{generation_id}.html")
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Business Flow Chart</title>
        <script src="https://cdnjs.cloudflare.com/ajax/libs/raphael/2.3.0/raphael.min.js"></script>
        <script src="https://cdnjs.cloudflare.com/ajax/libs/flowchart/1.15.0/flowchart.min.js"></script>
        <style>
            body {{ margin: 0; padding: 20px; background: white; }}
            #canvas {{ width: 100%; height: 100%; }}
        </style>
    </head>
    <body>
        <div id="canvas"></div>
        <script>
            var diagram = flowchart.parse(`{flowchart_code}`);
            diagram.drawSVG('canvas', {{
                'line-width': 2,
                'line-length': 50,
                'text-margin': 10,
                'font-size': 14,
                'font': 'normal',
                'font-family': 'Arial',
                'font-weight': 'normal',
                'font-color': 'black',
                'line-color': '#7f8c8d',
                'element-color': '#3498db',
                'fill': '#A5D8F3',
                'yes-text': 'Yes',
                'no-text': 'No',
                'arrow-end': 'block',
                'scale': 1,
                'symbols': {{
                    'start': {{ 'font-color': 'black', 'element-color': '#3498db', 'fill': '#A5D8F3' }},
                    'end': {{ 'font-color': 'black', 'element-color': '#3498db', 'fill': '#A5D8F3' }}
                }}
            }});
        </script>
    </body>
    </html>
    """

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    logger.debug(f"HTML file saved to: {html_path}")

    svg_path = os.path.join(output_dir, f"flowchart_{generation_id}.svg")
    png_path = os.path.join(output_dir, f"flowchart_{generation_id}.png")
    options = Options()
    options.headless = True
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')

    try:
        driver = webdriver.Chrome(options=options)
        driver.get(f"file://{os.path.abspath(html_path)}")
        driver.implicitly_wait(5)
        svg_element = driver.find_element(By.ID, 'canvas').get_attribute('innerHTML')
        if not svg_element.startswith('<svg'):
            svg_content = f'<svg xmlns="http://www.w3.org/2000/svg" width="1600" height="1200">{svg_element}</svg>'
        else:
            svg_content = svg_element
        with open(svg_path, 'w', encoding='utf-8') as f:
            f.write(svg_content)
        logger.debug(f"SVG flowchart saved to: {svg_path}")

        # Convert SVG to high-resolution PNG
        svg2png(url=svg_path, write_to=png_path, output_width=2400, output_height=1800)
        logger.debug(f"High-resolution PNG flowchart saved to: {png_path}")
    except Exception as e:
        logger.error(f"Error generating flowchart image: {e}")
        png_path = None
    finally:
        driver.quit()

    return html_path, png_path

def create_bpg_pdf(documentation, png_path, output_dir, generation_id):
    """Create a PDF document with Title, Description, Flowchart on page 1, Defects, Acceptance Criteria, and Appendix on page 2"""
    pdf_path = os.path.join(output_dir, f"business_process_guide_{generation_id}.pdf")
    doc = SimpleDocTemplate(pdf_path, pagesize=letter, leftMargin=0.75*inch, rightMargin=0.75*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    normal = styles['Normal']
    normal.fontSize = 10
    heading = styles['Heading2']
    heading.fontSize = 14
    heading.spaceAfter = 6
    title_style = styles['Heading1']
    title_style.fontSize = 18
    title_style.alignment = 1
    elements = []

    # Page 1: Title, Description, Flowchart
    elements.append(Paragraph(documentation.get("title", "Business Process Guide"), title_style))
    elements.append(Spacer(1, 0.2 * inch))

    elements.append(Paragraph("Description", heading))
    elements.append(Paragraph(documentation.get("description", "No description available"), normal))
    elements.append(Spacer(1, 0.1 * inch))

    if png_path and os.path.exists(png_path):
        elements.append(Paragraph("Flowchart", heading))
        img = Image(png_path, width=6.5*inch, height=4.5*inch)
        img.hAlign = 'CENTER'
        elements.append(img)

    elements.append(PageBreak())  # Force page break after Flowchart

    # Page 2: Common Defects, Acceptance Criteria, Appendix
    if documentation.get("common_defects"):
        elements.append(Paragraph("Common Defects", heading))
        headers = ["Defect", "Occurrence"]
        data = [headers]
        for defect in documentation["common_defects"]:
            data.append([Paragraph(str(defect.get("defect", "")), normal), Paragraph(str(defect.get("occurrence", "")), normal)])
        table = Table(data, colWidths=[3.5*inch, 3.5*inch])
        table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 0.2 * inch))

    if documentation.get("acceptance_criteria"):
        elements.append(Paragraph("Acceptance Criteria", heading))
        for crit in documentation["acceptance_criteria"]:
            elements.append(Paragraph(f"• {crit}", normal))
        elements.append(Spacer(1, 0.2 * inch))

    if documentation.get("appendix_table"):
        elements.append(Paragraph("Appendix", heading))
        headers = ["Parameter", "Value"]
        data = [headers]
        for param in documentation["appendix_table"]:
            data.append([Paragraph(str(param.get("parameter", "")), normal), Paragraph(str(param.get("value", "")), normal)])
        table = Table(data, colWidths=[3.5*inch, 3.5*inch])
        table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
        ]))
        elements.append(table)

    doc.build(elements)
    logger.debug(f"PDF saved to: {pdf_path}")
    return pdf_path

def create_bpg_word(documentation, png_path, output_dir, generation_id):
    """Create a Word document with Title, Description, Flowchart on page 1, Defects, Acceptance Criteria, and Appendix on page 2"""
    doc_dir = os.path.join(output_dir, 'documents')
    os.makedirs(doc_dir, exist_ok=True)
    doc_path = os.path.join(doc_dir, f"business_flow_{generation_id}.docx")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Page 1: Title, Description, Flowchart
    title = doc.add_heading(documentation.get("title", "Business Process Guide"), 0)
    title.paragraph_format.space_after = Pt(12)

    doc.add_heading("Description", 1)
    doc.add_paragraph(documentation.get("description", "No description available"))

    doc.add_heading("Flowchart", 1)
    if png_path and os.path.exists(png_path):
        try:
            img_para = doc.add_paragraph()
            img_para.paragraph_format.space_before = Pt(0)
            img_para.paragraph_format.space_after = Pt(6)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(png_path, width=Inches(5.5))
            logger.debug("Successfully added high-resolution PNG flowchart to document")
        except Exception as e:
            logger.error(f"Error adding PNG to document: {e}")
            doc.add_paragraph("Failed to embed flowchart image.")
    else:
        doc.add_paragraph("Flowchart image could not be generated.")

    doc.add_page_break()  # Force page break after Flowchart

    # Page 2: Common Defects, Acceptance Criteria, Appendix
    doc.add_heading("Common Defects", 1)
    if documentation.get("common_defects"):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        hdr = table.rows[0].cells
        hdr[0].text = "Defect"
        hdr[1].text = "Occurrence"
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True
        for d in documentation["common_defects"]:
            row = table.add_row().cells
            row[0].text = str(d.get("defect", ""))
            row[1].text = str(d.get("occurrence", ""))
    else:
        doc.add_paragraph("No common defects identified.")

    doc.add_heading("Acceptance Criteria", 1)
    if documentation.get("acceptance_criteria"):
        for crit in documentation["acceptance_criteria"]:
            doc.add_paragraph(f"• {crit}")
    else:
        doc.add_paragraph("No acceptance criteria provided.")

    doc.add_heading("Appendix", 1)
    if documentation.get("appendix_table"):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        hdr = table.rows[0].cells
        hdr[0].text = "Parameter"
        hdr[1].text = "Value"
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True
        for p in documentation["appendix_table"]:
            row = table.add_row().cells
            row[0].text = str(p.get("parameter", ""))
            row[1].text = str(p.get("value", ""))
    else:
        doc.add_paragraph("No appendix parameters identified.")

    doc.save(doc_path)
    logger.debug(f"Word document saved to: {doc_path}")
    return doc_path

def generate_and_save_bpg(user_story_data: dict):
    """Generate and save BPG documents with Title, Description, Flowchart on page 1, Defects, Acceptance Criteria, Appendix on page 2"""
    generation_id = str(uuid.uuid4())
    output_dir = 'output'
    os.makedirs(output_dir, exist_ok=True)

    documentation = generate_documentation(user_story_data)
    if not documentation:
        logger.error("Failed to generate documentation.")
        return None, None, None, None

    flow_definition = analyze_flow_structure(user_story_data)

    json_path = os.path.join(output_dir, f"flow_{generation_id}.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(flow_definition, f, indent=2)
    logger.debug(f"Flow definition saved to: {json_path}")

    flowchart_code, script_path, flowchart_path = json_to_pyflowchart(flow_definition, output_dir, generation_id)

    if flowchart_code:
        html_path, png_path = generate_flowchart_image(flowchart_code, output_dir, generation_id)
        pdf_path = create_bpg_pdf(documentation, png_path, output_dir, generation_id)
        doc_path = create_bpg_word(documentation, png_path, output_dir, generation_id)
        return doc_path, pdf_path, html_path, json_path
    else:
        logger.warning("Failed to generate flowchart code.")
        pdf_path = create_bpg_pdf(documentation, None, output_dir, generation_id)
        doc_path = create_bpg_word(documentation, None, output_dir, generation_id)
        return doc_path, pdf_path, None, json_path

# FastAPI endpoint
@app.post("/generate_bpg", response_model=BPGResponse)
async def generate_bpg_endpoint(request: UserStoryRequest):
    """Generate BPG documents via FastAPI"""
    logger.debug(f"Received raw request body: {request}")
    try:
        # Convert Pydantic model to dict
        user_story_data = request.dict()
        logger.debug(f"Processing user story data: {user_story_data}")
        doc_path, pdf_path, html_path, json_path = generate_and_save_bpg(user_story_data)
        if not doc_path or not pdf_path:
            raise HTTPException(status_code=500, detail="Failed to generate BPG documents")
        return BPGResponse(
            word_path=doc_path,
            pdf_path=pdf_path,
            html_path=html_path,
            json_path=json_path,
            status="success"
        )
    except Exception as e:
        logger.error(f"Error in FastAPI endpoint: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# MCP tool
@mcp.tool(description="Generate Business Process Guide (BPG) documents from a user story.")
async def generate_bpg_mcp(user_story_json: str) -> dict:
    logger.debug(f"Processing user story via MCP: {user_story_json}")
    try:
        user_story_data = json.loads(user_story_json)
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in MCP input: {e}")
        return {"status": "failed", "error": "Invalid JSON input"}
    doc_path, pdf_path, html_path, json_path = generate_and_save_bpg(user_story_data)
    status = "success" if doc_path and pdf_path else "failed"
    return {
        "word_path": doc_path,
        "pdf_path": pdf_path,
        "html_path": html_path,
        "json_path": json_path,
        "status": status
    }

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="BPG Generator: Generate Business Process Guides")
    parser.add_argument("--mode", choices=["cli", "fastapi"], default="cli", help="Run mode: 'cli' for command-line or 'fastapi' for API server")
    parser.add_argument("--user-story-json", type=str, help="User story JSON file or string for CLI mode")
    parser.add_argument("--host", type=str, default="0.0.0.0", help="Host for FastAPI server")
    parser.add_argument("--port", type=int, default=8000, help="Port for FastAPI server")
    args = parser.parse_args()

    if args.mode == "fastapi":
        logger.info(f"Starting FastAPI server on {args.host}:{args.port}")
        uvicorn.run(app, host=args.host, port=args.port)
    else:
        # CLI mode
        if args.user_story_json:
            try:
                if os.path.isfile(args.user_story_json):
                    with open(args.user_story_json, 'r', encoding='utf-8') as f:
                        user_story_data = json.load(f)
                else:
                    user_story_data = json.loads(args.user_story_json)
            except json.JSONDecodeError as e:
                print(f"Error: Invalid JSON input: {e}")
                sys.exit(1)
        else:
            print("Please enter your user story JSON (type 'done' on a new line when finished):")
            lines = []
            while True:
                line = input()
                if line.lower() == 'done':
                    break
                lines.append(line)
            user_story_json = ''.join(lines)
            if not user_story_json.strip():
                print("No user story provided. Exiting.")
                sys.exit(1)
            try:
                user_story_data = json.loads(user_story_json)
            except json.JSONDecodeError as e:
                print(f"Error: Invalid JSON input: {e}")
                sys.exit(1)

        print("\nProcessing user story...")
        doc_path, pdf_path, html_path, json_path = generate_and_save_bpg(user_story_data)
        print(f"\nDocument generation completed!")
        print(f"Word document path: {doc_path}")
        print(f"PDF document path: {pdf_path}")
        if html_path:
            print(f"Interactive flowchart path: {html_path}")
        print(f"Flow definition path: {json_path}")

        try:
            if os.name == 'nt':  # Windows
                os.startfile(doc_path)
                os.startfile(pdf_path)
                print("Documents opened automatically.")
                if html_path:
                    os.startfile(html_path)
                    print("Interactive flowchart opened automatically.")
            else:
                print("You can open the documents at the paths above.")
                if html_path:
                    print("Open the HTML file in a web browser to view the interactive flowchart.")
        except Exception as e:
            print(f"Could not open documents automatically: {e}")
            print("Please open the documents at the paths above.")